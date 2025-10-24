from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm
# ADDED THESE IMPORTS to fix NameError: name 'OxmlElement' is not defined
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv, sys, os, textwrap
from io import BytesIO

# --- Configuration ---
# NOTE: Ensure the following paths/files exist in your project directory
ICON_DIR = "icons" 
FONT_PATH = "arial.ttf" 
CSV_FILE = "seeds.csv"

# --- Read CSV data ---
data = {}
try:
    with open(CSV_FILE, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            data[row["SEED_NAME"].strip().upper()] = row
except FileNotFoundError:
    print(f"❌ Error: '{CSV_FILE}' not found. Please ensure the CSV file is in the correct directory.")
    sys.exit(1)
except Exception as e:
    print(f"❌ Error reading '{CSV_FILE}': {e}")
    sys.exit(1)

# --- Get input ---
seed_name = sys.argv[1].upper() if len(sys.argv) > 1 else "BOTTLE GOURD"
if seed_name not in data:
    print(f"❌ Seed '{seed_name}' not found in {CSV_FILE}. Available seeds: {list(data.keys())}")
    sys.exit(1)

info = data[seed_name]

# --- Card dimensions (300 DPI) ---
# 1 inch = 2.54 cm
cm_to_px = lambda cm_val: int((cm_val / 2.54) * 300)
CARD_WIDTH_CM, CARD_HEIGHT_CM = 3, 4.5
W, H = cm_to_px(CARD_WIDTH_CM), cm_to_px(CARD_HEIGHT_CM)

# --- Create seed card function ---
def create_seed_card(info_dict):
    # Image.new() is defined because of 'from PIL import Image' at the top
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    # Colors and fonts
    border_color = (0, 90, 0)
    border_width = 8
    text_color = (0, 0, 0)

    # Try loading custom font, fallback to default
    try:
        title_font_size = 48
        text_font_size = 34
        title_font = ImageFont.truetype(FONT_PATH, title_font_size)
        text_font = ImageFont.truetype(FONT_PATH, text_font_size)
    except IOError:
        # Fallback setup
        title_font = ImageFont.load_default()
        text_font = ImageFont.load_default()
        title_font_size = 32
        text_font_size = 20

    # --- Draw border (FIXED RIGHT SIDE CLIPPING) ---
    # Draw a rectangle slightly inset to ensure border is fully visible
    draw.rectangle(
        [
            (border_width // 2, border_width // 2), 
            (W - (border_width // 2) - 1, H - (border_width // 2) - 1)
        ], 
        outline=border_color, 
        width=border_width
    )

    # --- Margins and Spacing ---
    left_padding = 15
    right_padding = 15
    top_padding = 40
    icon_size = 40
    line_spacing = 10
    
    # --- Dynamic title ---
    seed_title = info_dict["SEED_NAME"].upper()
    current_title_font_size = title_font_size
    current_title_font = title_font

    # Adjust title font size to fit width
    max_title_width = W - left_padding - right_padding 
    while True:
        bbox = draw.textbbox((0, 0), seed_title, font=current_title_font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        if tw <= max_title_width or current_title_font_size <= 20:
            break
        current_title_font_size -= 2
        try:
            current_title_font = ImageFont.truetype(FONT_PATH, current_title_font_size)
        except IOError:
            current_title_font = ImageFont.load_default()

    # Center title horizontally
    y_current = top_padding
    x_title = (W - tw) / 2
    draw.text((x_title, y_current), seed_title, fill=border_color, font=current_title_font)
    y_current += th + 15

    # --- Icons Map ---
    icon_map = {
        "Position": os.path.join(ICON_DIR, "sun.png"),
        "Sow": os.path.join(ICON_DIR, "seedling.png"),
        "Germination": os.path.join(ICON_DIR, "hourglass.png"),
        "Spacing": os.path.join(ICON_DIR, "ruler.png"),
        "Height": os.path.join(ICON_DIR, "height.png"),
        "Water": os.path.join(ICON_DIR, "water.png"),
        "Note": os.path.join(ICON_DIR, "note.png"),
    }
    
    # --- Fields to display (FIXED MISSING WATER AND NOTE) ---
    fields_to_display = [
        ("Position", info_dict["POSITION"]),
        ("Sow", f"{info_dict['SOW_DEPTH_MM']} mm"),
        ("Germination", f"{info_dict['GERMINATION_DAYS']} days"),
        ("Spacing", f"{info_dict['SPACING_CM']} cm"),
        ("Height", f"{info_dict['MATURE_HEIGHT_M']} m"),
        ("Water", info_dict["WATER_NEEDS"]), # FIXED: WATER_NEEDS added
        ("Note", info_dict["SPECIAL_NOTE_1"]), # FIXED: SPECIAL_NOTE_1 added
    ]

    # Calculate text wrapping parameters
    text_start_x = left_padding + icon_size + 5 
    available_text_width = W - text_start_x - right_padding
    
    # Get text font metrics for line height
    ascent, descent = text_font.getmetrics()
    single_line_height = ascent + descent + 5
    
    # CRITICAL FIX: Adjusting average character width for generous wrapping
    representative_string = "MMWWiiilll" 
    text_length = draw.textlength(representative_string, font=text_font)
    
    # Reduce the calculated average width by 10% to make the line limit more generous
    avg_char_width = (text_length / len(representative_string)) * 0.9 
    wrap_width = max(10, int(available_text_width / avg_char_width))

    for label, value in fields_to_display:
        icon_path = icon_map.get(label)
        
        # Load and paste icon
        if icon_path and os.path.exists(icon_path):
            try:
                icon_img = Image.open(icon_path).convert("RGBA").resize((icon_size, icon_size))
                icon_y_align = y_current 
                img.paste(icon_img, (left_padding, int(icon_y_align)), icon_img)
            except Exception as e:
                # print(f"⚠️ Warning: Could not load icon '{icon_path}': {e}") # Debugging aid
                pass
        
        # Prepare text, (Note does not need 'Note: ' prefix)
        text_to_display = f"{label}: {value}" if label != "Note" else value 
        wrapped_lines = textwrap.wrap(text_to_display, width=wrap_width)
        
        # Draw each line of wrapped text
        for line in wrapped_lines:
            draw.text((text_start_x, y_current), line, fill=text_color, font=text_font)
            y_current += single_line_height 
        
        y_current += line_spacing

    return img

# --- Word document setup ---
doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.left_margin = Cm(0.5)
section.right_margin = Cm(0.5)
section.top_margin = Cm(1)
section.bottom_margin = Cm(1)

# Table for cards (5x5 grid)
rows, cols = 5, 5
table = doc.add_table(rows=rows, cols=cols)
table.autofit = False

# Set column/cell widths and remove cell padding for snug fit (FIXED BY ADDING IMPORTS)
for col in table.columns:
    col.width = Cm(CARD_WIDTH_CM)
    for cell in col.cells:
        cell.width = Cm(CARD_WIDTH_CM)
        
        # Remove internal cell margins/padding using OXML
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'bottom', 'left', 'right']:
            margin_element = OxmlElement(f'w:{margin_name}')
            margin_element.set(qn('w:w'), '0')
            margin_element.set(qn('w:type'), 'dxa')
            tcMar.append(margin_element)
        tcPr.append(tcMar)


# Insert cards into the table cells
for r_idx in range(rows):
    for c_idx in range(cols):
        cell = table.cell(r_idx, c_idx)
        card_img = create_seed_card(info)
        img_stream = BytesIO()
        card_img.save(img_stream, format="PNG")
        img_stream.seek(0)
        
        p = cell.paragraphs[0]
        p.clear() 
        run = p.add_run()
        run.add_picture(img_stream, width=Cm(CARD_WIDTH_CM), height=Cm(CARD_HEIGHT_CM))
        
        # Center the image within the cell
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out_name = f"{seed_name.replace(' ', '_')}_cards.docx"
doc.save(out_name)
print(f"✅ Word document created with {rows*cols} '{seed_name}' cards per page: {out_name}")




